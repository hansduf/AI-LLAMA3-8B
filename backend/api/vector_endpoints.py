"""
🚀 Vector API Endpoints
API endpoints for document embedding and vector search
"""

import logging
from typing import List, Dict, Any, Optional
from fastapi import HTTPException, UploadFile, File, Form
from pydantic import BaseModel
from datetime import datetime
import json
import uuid

logger = logging.getLogger(__name__)

# Pydantic models for API
class BulkUploadRequest(BaseModel):
    """Bulk document upload request"""
    process_immediately: bool = False
    extract_text_only: bool = True  # As requested - as text processing

class SearchRequest(BaseModel):
    """Vector search request"""
    query: str
    limit: Optional[int] = 10
    similarity_threshold: Optional[float] = 0.5
    document_ids: Optional[List[str]] = None

class SearchResponse(BaseModel):
    """Vector search response"""
    results: List[Dict[str, Any]]
    total_results: int
    query: str
    search_time_ms: int

class ProcessingStatusResponse(BaseModel):
    """Background processing status response"""
    task_id: str
    status: str
    document_id: str
    filename: str
    created_at: datetime
    progress_info: Dict[str, Any]

class VectorAPIRouter:
    """
    API Router for vector functionality
    Modular approach - can be imported into main.py easily
    """
    
    def __init__(self):
        self.endpoints_registered = False
    
    def register_endpoints(self, app):
        """
        Register vector API endpoints with FastAPI app
        Called from main.py to add endpoints
        """
        if self.endpoints_registered:
            logger.warning("Vector endpoints already registered")
            return
            
        from vector import (
            vector_migration, embedding_generator, text_chunker,
            vector_search, background_processor
        )
        from database.connection import SessionLocal
        
        def extract_text_from_file(filename: str, content: bytes) -> str:
            """Extract text from file content"""
            try:
                import io as file_io
                
                # Get file extension
                ext = filename.lower().split('.')[-1]
                
                if ext == 'pdf':
                    # Extract text from PDF
                    try:
                        from PyPDF2 import PdfReader
                        pdf_reader = PdfReader(file_io.BytesIO(content))
                        text_content = ""
                        for page in pdf_reader.pages:
                            text_content += page.extract_text() + "\n"
                        return text_content.strip()
                    except ImportError:
                        logger.warning("PyPDF2 not available, trying as text")
                        return content.decode('utf-8', errors='ignore')
                    
                elif ext in ['docx', 'doc']:
                    # Extract text from Word document
                    try:
                        from docx import Document
                        doc = Document(file_io.BytesIO(content))
                        text_content = ""
                        for paragraph in doc.paragraphs:
                            text_content += paragraph.text + "\n"
                        return text_content.strip()
                    except ImportError:
                        logger.warning("python-docx not available, trying as text")
                        return content.decode('utf-8', errors='ignore')
                    
                elif ext == 'txt':
                    # Plain text file
                    return content.decode('utf-8', errors='ignore')
                    
                else:
                    # Try to decode as text
                    return content.decode('utf-8', errors='ignore')
                    
            except Exception as e:
                logger.error(f"❌ Error extracting text from {filename}: {e}")
                return ""
        
        @app.post("/api/vector/setup")
        def setup_vector_database():
            """Setup vector database (run migration)"""
            try:
                result = vector_migration.run_migration()
                return {"message": "Vector database setup completed", "status": "success", "result": result}
            except Exception as e:
                logger.error(f"❌ Vector setup failed: {e}")
                raise HTTPException(status_code=500, detail=f"Setup failed: {str(e)}")
        
        @app.post("/api/vector/upload-bulk")
        def upload_documents_bulk(
            files: List[UploadFile] = File(...),
            request_data: str = Form(...)
        ):
            """
            Bulk document upload with background embedding processing
            As requested: mass document upload with background processing
            """
            try:
                # Parse request data
                try:
                    upload_request = BulkUploadRequest.parse_raw(request_data)
                except Exception as e:
                    raise HTTPException(status_code=400, detail=f"Invalid request data: {e}")
                
                if not files:
                    raise HTTPException(status_code=400, detail="No files provided")
                
                logger.info(f"📚 Bulk upload started: {len(files)} files")
                
                # Start background processor if not running
                if not background_processor.is_running:
                    import asyncio
                    asyncio.create_task(background_processor.start_background_worker())
                
                uploaded_documents = []
                processing_tasks = []
                
                db = SessionLocal()
                try:
                    for file in files:
                        # Validate file
                        if not file.filename:
                            continue
                            
                        # Read file content
                        content = file.file.read()
                        
                        # Extract text as requested (as text processing)
                        text_content = extract_text_from_file(file.filename, content)
                        
                        if not text_content:
                            logger.warning(f"⚠️ No text extracted from {file.filename}")
                            continue
                        
                        # Save document to database
                        document_id = self._save_document_to_db(
                            db=db,
                            filename=file.filename,
                            content=text_content,
                            file_size=len(content)
                        )
                        
                        uploaded_documents.append({
                            "document_id": document_id,
                            "filename": file.filename,
                            "size": len(content),
                            "text_length": len(text_content)
                        })
                        
                        # Queue for background processing
                        import asyncio
                        task_id = asyncio.run(background_processor.queue_document_for_processing(
                            document_id=document_id,
                            filename=file.filename
                        ))
                        processing_tasks.append(task_id)
                    
                    db.commit()
                    
                    return {
                        "message": f"Bulk upload completed: {len(uploaded_documents)} documents",
                        "documents": uploaded_documents,
                        "processing_tasks": processing_tasks,
                        "background_processing": True
                    }
                    
                finally:
                    db.close()
                
            except Exception as e:
                logger.error(f"❌ Bulk upload failed: {e}")
                raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")
        
        @app.post("/api/vector/upload")
        def upload_single_document(file: UploadFile = File(...)):
            """
            Single document upload with immediate processing
            """
            try:
                if not file.filename:
                    raise HTTPException(status_code=400, detail="No filename provided")
                
                logger.info(f"📄 Single upload started: {file.filename}")
                
                # Read file content
                content = file.file.read()
                
                # Extract text content
                text_content = extract_text_from_file(file.filename, content)
                
                if not text_content or len(text_content.strip()) < 10:
                    raise HTTPException(status_code=400, detail="Could not extract meaningful text from file")
                
                db = SessionLocal()
                try:
                    # Store document in database
                    document_id = str(uuid.uuid4())
                    
                    # Insert document
                    db.execute(text("""
                        INSERT INTO documents (id, filename, content, file_size, file_type, created_at, updated_at)
                        VALUES (%s, %s, %s, %s, %s, %s, %s)
                    """), [
                        document_id,
                        file.filename,
                        text_content,
                        len(content),
                        file.content_type or 'application/octet-stream',
                        datetime.now(),
                        datetime.now()
                    ])
                    
                    # Chunk the document
                    chunks = text_chunker.chunk_text(text_content, {"document_id": document_id})
                    
                    # Generate embeddings and store chunks
                    import asyncio
                    
                    for chunk in chunks:
                        # Generate embedding
                        embedding = asyncio.run(embedding_generator.generate_document_embedding(chunk.content))
                        
                        if embedding:
                            # Store chunk with embedding
                            db.execute(text("""
                                INSERT INTO document_chunks (chunk_id, document_id, content, chunk_index, metadata, embedding, created_at)
                                VALUES (%s, %s, %s, %s, %s, %s, %s)
                            """), [
                                chunk.chunk_id,
                                document_id,
                                chunk.content,
                                chunk.chunk_index,
                                json.dumps(chunk.metadata),
                                embedding,
                                datetime.now()
                            ])
                    
                    db.commit()
                    
                    return {
                        "message": "Document uploaded and processed successfully",
                        "document_id": document_id,
                        "filename": file.filename,
                        "chunks_created": len(chunks),
                        "status": "success"
                    }
                    
                finally:
                    db.close()
                
            except Exception as e:
                logger.error(f"❌ Single upload failed: {e}")
                raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")

        @app.post("/api/vector/search", response_model=SearchResponse)
        def search_documents(request: SearchRequest):
            """
            Vector search for document Q&A
            As requested: AI Q&A with document references
            """
            try:
                start_time = datetime.now()
                
                db = SessionLocal()
                try:
                    # Perform vector search
                    import asyncio
                    results = asyncio.run(vector_search.search_with_query_text(
                        query_text=request.query,
                        embedding_generator=embedding_generator,
                        session=db,
                        limit=request.limit,
                        similarity_threshold=request.similarity_threshold,
                        document_ids=request.document_ids
                    ))
                    
                    # Format results
                    formatted_results = []
                    for result in results:
                        formatted_results.append({
                            "chunk_id": result.chunk_id,
                            "document_id": result.document_id,
                            "document_name": result.document_name,
                            "content": result.content,
                            "similarity_score": result.similarity_score,
                            "chunk_index": result.chunk_index,
                            "metadata": result.metadata
                        })
                    
                    search_time = (datetime.now() - start_time).total_seconds() * 1000
                    
                    return SearchResponse(
                        results=formatted_results,
                        total_results=len(formatted_results),
                        query=request.query,
                        search_time_ms=int(search_time)
                    )
                    
                finally:
                    db.close()
                    
            except Exception as e:
                logger.error(f"❌ Vector search failed: {e}")
                raise HTTPException(status_code=500, detail=f"Search failed: {str(e)}")
        
        @app.get("/api/vector/documents")
        def search_documents_by_content(query: str, limit: int = 5):
            """Search documents by content similarity"""
            try:
                db = SessionLocal()
                try:
                    import asyncio
                    results = asyncio.run(vector_search.search_documents_by_content(
                        query_text=query,
                        embedding_generator=embedding_generator,
                        session=db,
                        limit=limit
                    ))
                    
                    return {
                        "query": query,
                        "documents": results,
                        "total_found": len(results)
                    }
                    
                finally:
                    db.close()
                    
            except Exception as e:
                logger.error(f"❌ Document search failed: {e}")
                raise HTTPException(status_code=500, detail=f"Document search failed: {str(e)}")
        
        @app.get("/api/vector/processing-status")
        def get_processing_status():
            """Get background processing status"""
            try:
                stats = background_processor.get_processing_stats()
                return {
                    "processing_stats": stats,
                    "timestamp": datetime.now().isoformat()
                }
            except Exception as e:
                logger.error(f"❌ Error getting processing status: {e}")
                raise HTTPException(status_code=500, detail=f"Status check failed: {str(e)}")
        
        @app.get("/api/vector/task-status/{task_id}")
        def get_task_status(task_id: str):
            """Get specific task status"""
            try:
                task = background_processor.get_task_status(task_id)
                if not task:
                    raise HTTPException(status_code=404, detail="Task not found")
                
                return {
                    "task_id": task.task_id,
                    "status": task.status.value,
                    "document_id": task.document_id,
                    "filename": task.filename,
                    "created_at": task.created_at.isoformat(),
                    "started_at": task.started_at.isoformat() if task.started_at else None,
                    "completed_at": task.completed_at.isoformat() if task.completed_at else None,
                    "error_message": task.error_message,
                    "retry_count": task.retry_count
                }
                
            except HTTPException:
                raise
            except Exception as e:
                logger.error(f"❌ Error getting task status: {e}")
                raise HTTPException(status_code=500, detail=f"Task status check failed: {str(e)}")
        
        @app.post("/api/vector/start-processor")
        def start_background_processor():
            """Start background embedding processor"""
            try:
                import asyncio
                asyncio.create_task(background_processor.start_background_worker())
                return {"message": "Background processor started", "status": "running"}
            except Exception as e:
                logger.error(f"❌ Error starting processor: {e}")
                raise HTTPException(status_code=500, detail=f"Failed to start processor: {str(e)}")
        
        @app.post("/api/vector/stop-processor")
        def stop_background_processor():
            """Stop background embedding processor"""
            try:
                import asyncio
                asyncio.run(background_processor.stop_background_worker())
                return {"message": "Background processor stopped", "status": "stopped"}
            except Exception as e:
                logger.error(f"❌ Error stopping processor: {e}")
                raise HTTPException(status_code=500, detail=f"Failed to stop processor: {str(e)}")
        
        self.endpoints_registered = True
        logger.info("✅ Vector API endpoints registered successfully")
    
    def _extract_text_from_file(self, filename: str, content: bytes) -> str:
        """
        Extract text from file content (as text processing as requested)
        """
        try:
            from PyPDF2 import PdfReader
            from docx import Document
            import io
            
            # Get file extension
            ext = filename.lower().split('.')[-1]
            
            if ext == 'pdf':
                # Extract text from PDF
                pdf_reader = PdfReader(io.BytesIO(content))
                text_content = ""
                for page in pdf_reader.pages:
                    text_content += page.extract_text() + "\n"
                return text_content.strip()
                
            elif ext in ['docx', 'doc']:
                # Extract text from Word document
                doc = Document(io.BytesIO(content))
                text_content = ""
                for paragraph in doc.paragraphs:
                    text_content += paragraph.text + "\n"
                
                # Also extract table content as text (as requested)
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            row_text.append(cell.text.strip())
                        text_content += " | ".join(row_text) + "\n"
                
                return text_content.strip()
                
            elif ext == 'txt':
                # Plain text file
                return content.decode('utf-8', errors='ignore')
                
            else:
                # Try to decode as text
                return content.decode('utf-8', errors='ignore')
                
        except Exception as e:
            logger.error(f"❌ Error extracting text from {filename}: {e}")
            return ""
    
    def _save_document_to_db(self, db, filename: str, content: str, file_size: int) -> str:
        """Save document to database"""
        try:
            from sqlalchemy import text
            
            document_id = str(uuid.uuid4())
            
            # Insert document
            insert_query = """
                INSERT INTO documents 
                (id, filename, content, file_size, embeddings_completed, embedding_status, created_at, updated_at)
                VALUES (%s, %s, %s, %s, %s, %s, NOW(), NOW())
            """
            
            db.execute(text(insert_query), [
                document_id,
                filename,
                content,
                file_size,
                False,  # embeddings_completed
                'pending'  # embedding_status
            ])
            
            logger.info(f"💾 Document saved: {filename} (id: {document_id})")
            return document_id
            
        except Exception as e:
            logger.error(f"❌ Error saving document {filename}: {e}")
            raise

# Global router instance
vector_api_router = VectorAPIRouter()
